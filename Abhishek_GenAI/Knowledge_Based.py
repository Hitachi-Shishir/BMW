import boto3
import json
from botocore.exceptions import ClientError
from docx import Document
from docx.shared import Pt
import os
import pyodbc
import sys
from datetime import datetime


DatabaseString = sys.argv[2]
modelId = "meta.llama3-70b-instruct-v1:0"
# Initialize the Bedrock client (assuming your AWS credentials are correctly configured in your environment)
CONNECTION_STRING = (
    'DRIVER={ODBC Driver 17 for SQL Server};'
    'SERVER=EC2AMAZ-G25T1EO;'
    'DATABASE='+DatabaseString+';'
    'UID=sa;'
    'PWD=Kaseya@123;'
)

def Database_connectivity(Ticket_ID):
    # Connect to the database
    try:
        #print("database")
        with pyodbc.connect(CONNECTION_STRING) as connection:
            cursor = connection.cursor()
            query = "select ticketsummary,TicketDesc,solutionNote from SDIncident where ticketref=?"
            #print("before execution")
            cursor.execute(query, (Ticket_ID,))
            #print("after execution")
            row = cursor.fetchall()
            issue_description = row[0][1]
            title = row[0][0]
            resolution_steps = row[0][2]
  
            query2 = "select NoteDesc from SD_TicketNotes where TicketID=?"
            cursor.execute(query, (Ticket_ID,))
            row2 = cursor.fetchall()
            additional_notes = row2[0]
         
            cursor.close()
            #print("after execution")
            generate_knowledge_base_with_bedrock (Ticket_ID, issue_description,title, resolution_steps, additional_notes)
            
            #combined_string = "\n".join(" ".join(item) for item in row)
            #if combined_string:
                #print (combined_string,"tytdsyugfsdgyufg")
            #    return combined_string
            #else:
                #return("No data to summarize")
                #sys.exit()
    except Exception as e:
        #print(f"Database error: {e}")
        return None

def generate_knowledge_base_with_bedrock(ticket_id, issue_description, title, resolution_steps, additional_notes):
    client = boto3.client("bedrock-runtime", region_name="us-east-1", aws_access_key_id="AKIAWF7TD545IFJKEZOB", aws_secret_access_key="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")

    # Create the input prompt for the model
    prompt_text = f"""
    You are a technical support expert tasked with creating a formal knowledge base document for a service desk ticket.
    Please create a well-structured and detailed knowledge base entry based on the following information. 
    The document should be formatted professionally, using concise language and proper headings:

    Ticket ID: {ticket_id}
    Title: {title}

    Issue Description:
    {issue_description}

    Resolution Steps:
    {resolution_steps}

    Additional Notes:
    {additional_notes}

    The knowledge base article should be structured in the following sections:
    - **Overview**: A brief summary of the issue.
    - **Detailed Issue Description**: Explanation of the problem as reported.
    - **Step-by-Step Resolution**: A clear list of steps taken to resolve the issue.
    - **Additional Information**: Any extra details or notes that might be useful.
    - **Preventive Measures**: Suggestions to avoid similar issues in the future.

    Format this as a professional knowledge base article with appropriate headings and formal language.
    """

    # Create the JSON body to send to the Bedrock model with the correct key 'prompt'
    input_data = {
        "prompt": prompt_text,
        "max_gen_len": 512,
        "temperature": 0.5,
        "top_p": 0.9
    }

    try:
        # Call the Bedrock model using the invoke_model method
        response = client.invoke_model(
            modelId=modelId,  # Replace with the actual model ID you want to use
            contentType="application/json",
            accept="application/json",
            body=json.dumps(input_data).encode('utf-8')
        )

        # Read the response from the Bedrock model
        response_body = response['body'].read().decode('utf-8')
        #print("Response from Bedrock model:", response_body)  # Debugging line to print the response

        # Parse the response to extract the generated content
        result = json.loads(response_body)
        

        # Check the structure of the response and extract the text data accordingly
        generated_content = result.get('generation', '')
        print(generated_content)

        # Define the file name based on the ticket ID
        location = "D:\\ServiceDesk30_1\\ServiceDeskV2\\KBARTICLE\\"
        #location = "D:\\" + DatabaseString + "\\KBARTICLE\\"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"{location}KB_{ticket_id}_{timestamp}.docx"

        # Create a new Document
        doc = Document()
        
        # Add content to the document
        doc.add_heading(f'Knowledge Base Article: {title}', level=1)
        doc.add_paragraph(f'**Ticket ID**: {ticket_id}\n')
        doc.add_page_break()
        # Assuming the content is structured properly, we will split it into sections.
        sections = generated_content.split('\n\n')  # Split by double new lines for sections

        for section in sections:
            # Add each section as a heading and content
            if section.strip():
                # Split the section into title and body
                lines = section.split('\n')
                doc.add_heading(lines[0], level=2).alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph = doc.add_paragraph("\n".join(lines[1:]).strip())
                paragraph.style.font.size = Pt(12)



                # Add a page break to start the body from the second page
                
                # Add body

                


        # Save the document
        doc.save(file_name)
        #return(file_name)

        print(f"$$$ '{file_name}' $$$")
    
    except ClientError as e:
        print(f"ClientError: An error occurred while communicating with AWS Bedrock: {e}")
    except Exception as e:
        print(f"An error occurred while generating the knowledge base document: {e}")



if __name__ == "__main__":
    #print("intisalizing python")
    TicketID = sys.argv[1]
    Database_connectivity(TicketID)
    #print ("In main function",TicketID)
   
    # Example usage:
     #generate_knowledge_base_with_bedrock(ticket_id, title, issue_description, resolution_steps, root_cause, additional_notes)
