import json
import boto3
import sqlite3
import uuid
from datetime import datetime
import re
import pyodbc
from docx import Document
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

# Initialize AWS Bedrock client

client = boto3.client("bedrock-runtime", region_name="us-east-1", aws_access_key_id="AKIAWF7TD545IFJKEZOB", aws_secret_access_key="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")
model_id = "meta.llama3-70b-instruct-v1:0"

DatabaseString = sys.argv[1]
CONNECTION_STRING = (
    'DRIVER={ODBC Driver 17 for SQL Server};'
    'SERVER=EC2AMAZ-G25T1EO;'
    'DATABASE='+DatabaseString+';'
    'UID=sa;'
    'PWD=Kaseya@123;'
)

def extract_sql_query(generative_output):
    # Use a regular expression to find the SQL query in the expected format
    sql_pattern = re.compile(r'SQLquery:\s*(.*?)\s*(?=YSQLquery:|$)', re.DOTALL | re.IGNORECASE)
    match = sql_pattern.search(generative_output)
    if match:
        return match.group(1)  # Return the SQL query without the surrounding text
    else:
        print("Please try with better prompt")
        return None
# Function to use AI model to convert user input into an SQL query
def convert_input_to_query_via_ai(user_input):
    try:
        # Define the prompt that will be sent to the AI model
        prompt = f""" f"Convert the following natural language request into an SQL query: \"{user_input}\". " 
                     f"The SQL database has a table 'vSDGenAIBased' with the following columns: 'TicketNumber', 'OrgName', 'ReqType', 'Stage',	'Status', 'Priority', 'Severity', 'Category', 'Resolution', 'ResolutionDescription', 'solutionNote', 'sourceType', TicketSummary', 'TicketDescription', 'submitterType', 'submitterName', 'submitterEmailAddress', 'submitterPhone', 'location', 'Department', 'Branch', 'TechnicianAssigned', 'CreationTime', 'DUEDATE', 'ResolutionTime', 'ResponseTime', 'TicketResolvedDate'. "
                     f"Generate a SQL query for this structure."
                     f"Output only the SQL code without any additional text, comments, or explanations."
                     f"The output should be formatted as: SQLquery:"<SQLquery>" "
                     f"Replace YSQLquery with the actual SQL query in single line without newline"
                     
        """
        input_data = {
        "prompt": prompt,
        "max_gen_len": 512,
        "temperature": 0.5,
        "top_p": 0.9
    }
        # Call AWS Bedrock to generate the SQL query
        response = client.invoke_model(
            modelId=model_id,  # Replace with your Bedrock model ID
            contentType="application/json",
            accept="application/json",
            body=json.dumps(input_data).encode('utf-8')  # JSON formatted prompt for the model
        )

        # Read the response from the Bedrock model
        response_body = response['body'].read().decode('utf-8')

        # Parse the response to extract the generated SQL query
        result = json.loads(response_body)
        #print(result)

        # Assuming 'generation' contains the SQL query or any other key from the model's output
        generative_output = result.get('generation', '').strip()
        
        sql_query = extract_sql_query(generative_output)
        
        if sql_query:
            #print(f"Generated SQL Query:", sql_query)
            return sql_query
            
                
        else:
            #print("I didn't able to understand, Please try again with better content")
            return None

    except Exception as e:
        #print(f"Error generating SQL query using AI model: {str(e)}")
        return None

# Function to execute the generated SQL query
def execute_query(query):
    try:
        # Connect to SQLite database
        #query = "`"+query+"`"
        with pyodbc.connect(CONNECTION_STRING) as connection:
            cursor = connection.cursor()
            query = cursor.execute(query)
            result = cursor.fetchall()
            #print(len(result))
            #print(f"Fetched {len(result)} rows.")
            column_names = [desc[0] for desc in cursor.description]
            #print("Column names:", len(column_names))
            
            # Check if the result is structured correctly
            if len(result) > 0:
                # Convert result to a DataFrame
                df = pd.DataFrame.from_records(result, columns=column_names)
                #print(f"DataFrame created with {df.shape[0]} rows and {df.shape[1]} columns.")
            else:
                #print("No data found")
                return None

            cursor.close()
        #print(df)      
        return (df)

    except Exception as e:
        print(f"Error executing query: {str(e)}")
        return None
def get_generative_ai_insight(prompt):
    # Replace with your OpenAI API key and model
    input_data = {
        "prompt": prompt
    }
    response = client.invoke_model(
            modelId="meta.llama2-13b-chat-v1",  # Replace with the actual model ID you want to use
            contentType="application/json",
            accept="application/json",
            body=json.dumps(input_data).encode('utf-8')
        )

        # Extract the model response
    result = json.loads(response['body'].read().decode('utf-8'))
        
    generated_text = result.get('generation', '')
    
    return (generated_text)
# Function to generate a report using AWS Bedrock
def generate_report(data):
    try:
        # Assume data is a DataFrame
        data_frame = data
        
        # Prepare the summary from the data
        total_tickets = len(data_frame)
        closed_tickets = len(data_frame[data_frame['Status'] == 'Closed'])
        avg_resolution_time = data_frame['ResolutionTime'].mean()
        
        # Create a breakdown of ticket types and counts
        ticket_counts = data_frame['ReqType'].value_counts()
        
        # Create a new Word Document
        doc = Document()
        
        # Add Title
        doc.add_heading('Service Desk Report', level=1)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=2)
        executive_summary = (
            f"This report presents an analysis of the service desk's performance. "
            f"Total Tickets: {total_tickets}, Closed Tickets: {closed_tickets}, "
            f"Average Resolution Time: {avg_resolution_time:.2f} hours."
        )
        doc.add_paragraph(executive_summary)
        
        # Key Performance Indicators (KPIs)
        doc.add_heading('Key Performance Indicators (KPIs)', level=2)
        doc.add_paragraph(f"- Total Tickets: {total_tickets}")
        doc.add_paragraph(f"- Closed Tickets: {closed_tickets}")
        doc.add_paragraph(f"- Average Resolution Time: {avg_resolution_time:.2f} hours")
        
        # Generative AI Insight for KPIs
        kpi_insight_prompt = (
            f"Analyze the following KPIs for a service desk: "
            f"Total Tickets: {total_tickets}, Closed Tickets: {closed_tickets}, "
            f"Average Resolution Time: {avg_resolution_time:.2f} hours. "
            f"Provide insights on performance and potential improvements."
        )
        kpi_insight = get_generative_ai_insight(kpi_insight_prompt)
        doc.add_heading('Insights from KPIs', level=2)
        doc.add_paragraph(kpi_insight)

        # Ticket Distribution Table
        doc.add_heading('Ticket Distribution Table', level=2)
        distribution_table = doc.add_table(rows=1, cols=2)
        distribution_table.style = 'Table Grid'
        hdr_cells = distribution_table.rows[0].cells
        hdr_cells[0].text = 'Ticket Type'
        hdr_cells[1].text = 'Count'
        
        for ticket_type, count in ticket_counts.items():
            row_cells = distribution_table.add_row().cells
            row_cells[0].text = ticket_type
            row_cells[1].text = str(count)

        # Departmental Resolution Time Comparison Table
        departmental_avg_time = data_frame.groupby('Department')['ResolutionTime'].mean().reset_index()
        
        doc.add_heading('Departmental Resolution Time Comparison Table', level=2)
        comparison_table = doc.add_table(rows=1, cols=2)
        comparison_table.style = 'Table Grid'
        hdr_cells = comparison_table.rows[0].cells
        hdr_cells[0].text = 'Department'
        hdr_cells[1].text = 'Avg. Resolution Time (hrs)'

        for index, row in departmental_avg_time.iterrows():
            row_cells = comparison_table.add_row().cells
            row_cells[0].text = row['Department']
            row_cells[1].text = f"{row['ResolutionTime']:.2f}"

        # Generative AI Insight for Departmental Comparison
        departmental_insight_prompt = (
            f"Analyze the average resolution times by department from the following data: "
            f"{departmental_avg_time.to_dict(orient='records')}. "
            f"Provide insights and potential areas of improvement."
        )
        departmental_insight = get_generative_ai_insight(departmental_insight_prompt)
        doc.add_heading('Insights from Departmental Resolution Times', level=2)
        doc.add_paragraph(departmental_insight)

        # Ticket Summary by Priority
        doc.add_heading('Ticket Summary by Priority', level=2)
        priority_summary = data_frame.groupby('Priority').agg(Open_Tickets=('Status', 'count')).reset_index()
        priority_table = doc.add_table(rows=1, cols=3)
        priority_table.style = 'Table Grid'
        hdr_cells = priority_table.rows[0].cells
        hdr_cells[0].text = 'Priority'
        hdr_cells[1].text = 'Open Tickets'
        hdr_cells[2].text = 'Closed Tickets'

        for _, row in priority_summary.iterrows():
            row_cells = priority_table.add_row().cells
            row_cells[0].text = row['Priority']
            row_cells[1].text = str(row['Open_Tickets'])
            closed_tickets = len(data_frame[(data_frame['Priority'] == row['Priority']) & (data_frame['Status'] == 'Closed')])
            row_cells[2].text = str(closed_tickets)

        # Trends and Observations
        doc.add_heading('Trends and Observations', level=2)
        trends_prompt = "Based on the ticket data provided, highlight key trends and observations that can be drawn from the ticket distribution, departmental performance, and priority analysis."
        trends_observations = get_generative_ai_insight(trends_prompt)
        doc.add_paragraph(trends_observations)

        # Conclusion and Recommendations
        doc.add_heading('Conclusion and Recommendations', level=2)
        recommendations_prompt = "Based on the insights from the service desk performance data, what actionable recommendations can be provided to enhance the efficiency of the service desk?"
        recommendations = get_generative_ai_insight(recommendations_prompt)
        doc.add_paragraph(recommendations)


        
        
        # Create a report structure

        #print(doc)
        location = "D:\\ServiceDesk30_1\\ServiceDeskV2\\AIReports\\"
        #location = "D:\\" + DatabaseString + "\\GenAIReports\\"
        # Save the document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = f"{location}_{timestamp}Service_Desk_Report.docx"
        doc.save(report_file)
        #print(f'Report generated successfully: {report_file}')
        
        return report_file  # Return the file path

    except Exception as e:
        #print(f"An error occurred: {e}")
        return None  # Return None or handle the error as needed




# Function to save the report locally
def save_report_locally(report_content, user_input):
    # Generate a unique report ID
    report_id = str(uuid.uuid4())
    
    # Create a report structure
    report_data = {
        'report_id': report_id,
        'user_input': user_input,
        'report_content': report_content,
        'timestamp': datetime.utcnow().isoformat()
    }

    # Store the report as a JSON file locally
    report_filename = f"report_{report_id}.docx"
    
 
    
    print(f"Report saved as {report_filename}")

    return report_filename

# Function to save report in the database
def save_report_to_db(db_path, report_content, user_input):
    try:
        # Connect to the SQLite database
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Generate a unique report ID
        report_id = str(uuid.uuid4())
        timestamp = datetime.utcnow().isoformat()

        # Insert the report into the database
        cursor.execute(
            "INSERT INTO reports (report_id, user_input, report_content, timestamp) VALUES (?, ?, ?, ?)",
            (report_id, user_input, report_content, timestamp)
        )
        
        # Commit the transaction and close the connection
        conn.commit()
        conn.close()
        
        #print(f"Report saved to database with report_id: {report_id}")
    
    except Exception as e:
        print(f"Error saving report to database: {str(e)}")

# Main logic to handle the flow
def print_doc_content(file_path):
    """Reads and prints the content of the Word document in a structured format."""
    doc = Document(file_path)
    
    print("\n--- Document Content ---")
    
    for para in doc.paragraphs:
        # Print headings
        if para.style.name.startswith('Heading'):
            heading_level = para.style.name[-1]  # Get heading level
            print(f"{'#' * int(heading_level)} {para.text}")
        else:
            print(para.text)

    # Loop through each table and print its content
    for table in doc.tables:
        print("\n--- Table ---")
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            print(" | ".join(row_data))
        print("--- End of Table ---")



if __name__ == "__main__":
    # Database path for storing tickets and reports
    #db_path = "ticket_data.db"
    
    # Example user input for generating the report
    user_input = sys.argv[2]

    # Step 1: Convert user input into SQL query using AI model
    query = convert_input_to_query_via_ai(user_input)

    if query:
        # Step 2: Execute the generated SQL query
        data = execute_query(query)
        

        if isinstance(data, pd.DataFrame) and not data.empty:
            # Step 3: Generate report using Bedrock model
            
            report_content = generate_report(data)
            print_doc_content(report_content)

            #if report_content:
                # Step 4a: Save the generated report locally
                #report_filename = save_report_locally(report_content, user_input)
                #rint(f"Report generated successfully and saved at {report_filename}")
                
                # Step 4b: Save the generated report in the database
                #save_report_to_db(report_content, user_input)
            
        else:
            print("No data found based on your request.")
    #else:
    #po0i    print("Could not process the user input into a valid query.")
